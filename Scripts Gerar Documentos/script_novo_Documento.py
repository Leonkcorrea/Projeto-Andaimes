import os
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm
import sqlite3

dbContato = sqlite3.connect('databaseContratos.db')
cContato = dbContato.cursor()

dbDados = sqlite3.connect('databaseDados.db')
cDados = dbDados.cursor()


def BUSCAR_CLIENTE(cpf):
    print(".")



def ADICIONAR_PECA_NOME(codigo):
    nome = 0
    return nome



def RETORNAR_VALOR(valor, genero, codigo): #Código da peça para buscar seu gênero
    valor_extenso = 0
    return valor_extenso



def RETORNAR_VALOR_FINANCEIRO(valor):
    valor_extenso = 0
    return valor_extenso


def ADD_DATABASE():
    connection = sqlite3.connect()


def add_space(qtd):
    s = ""
    i = 0
    while i < qtd:
        s = s + " "
        i += 1
    return s


def GERAR_DOCUMENTO():

    # Defenindo variáveis #
    document = docx.Document()
    styles = document.styles
    sections = document.sections

    # Configurando a página # -> # Convertendo inch para cm #
    for section in sections:
        section.page_width = Mm(210)
        section.page_height = Mm(310)
        section.top_margin = Inches(2.30 / 2.54)
        section.left_margin = Inches(1.50 / 2.54)
        section.right_margin = Inches(0.32 / 2.54)
        section.bottom_margin = Inches(3.49 / 2.54)

    # Definindo estilos # -> # [fontName_fontSize_bold/italic/underline] #
    style = []
    tnr_10_000 = styles.add_style('TNR_10_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_10_000)
    tnr_12_100 = styles.add_style('TNR_12_100', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_12_100)
    tnr_12_101 = styles.add_style('TNR_12_101', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_12_101)
    tnr_14_000 = styles.add_style('TNR_14_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_14_000)
    tnr_14_100 = styles.add_style('TNR_14_100', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_14_100)
    tnr_14_101 = styles.add_style('TNR_14_101', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_14_101)
    tnr_36_000 = styles.add_style('TNR_36_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_36_000)
    tnr_38_000 = styles.add_style('TNR_38_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_38_000)
    tnr_41_000 = styles.add_style('TNR_41_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_41_000)
    tnr_45_000 = styles.add_style('TNR_45_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_45_000)
    tnr_46_000 = styles.add_style('TNR_46_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_46_000)
    tnr_433_000 = styles.add_style('TNR_433_000', WD_STYLE_TYPE.PARAGRAPH); style.append(tnr_433_000)

    # Configurando estilos #
    for s in style:
        s.font.name = 'Times New Roman'
        s.font.size = Pt(int(s.name[4:-4]))
        if(s.name[-3] == '1'): s.font.bold = True
        if(s.name[-2] == '1'): s.font.italic = True
        if(s.name[-1] == '1'): s.font.underline = True

    # Escrevendo no documento #

    paragraphs = []

    paragraph = document.add_paragraph('', tnr_46_000); paragraphs.append(paragraph)
    
    paragraph = document.add_paragraph('(<carteira>)  Carteira', tnr_10_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('(<banco>)  Banco', tnr_10_000)
    run = paragraph.add_run(str(add_space(50)))
    run.font.size = Pt(14)
    run = paragraph.add_run('<numero>/<ano>')
    run.font.size = Pt(14)
    paragraphs.append(paragraph)
    
    paragraph = document.add_paragraph('', tnr_45_000); paragraphs.append(paragraph)
    
    paragraph = document.add_paragraph(('<nome>').upper(), tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<endereço> - <bairro>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<cep> - <cidade> - <estado>   Fone: <contato>', tnr_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_38_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<txt>', tnr_14_101)
    run = paragraph.add_run(' : ' + '<texto>')
    run.font.size = Pt(14)
    run.font.underline = False
    paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_46_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<prod1> - ( <qtd1> ) <quantidade1> <produto1>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod2> - ( <qtd2> ) <quantidade2> <produto2>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod3> - ( <qtd3> ) <quantidade3> <produto3>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod4> - ( <qtd4> ) <quantidade4> <produto4>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod5> - ( <qtd5> ) <quantidade5> <produto5>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod6> - ( <qtd6> ) <quantidade6> <produto6>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('<prod7> - ( <qtd7> ) <quantidade7> <produto7>', tnr_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph(str(add_space(112)) + '<datainicio>', tnr_14_000); paragraphs.append(paragraph)
    paragraph = document.add_paragraph(str(add_space(56)) + '<data>', tnr_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_41_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph(str(add_space(79)) + '<vlr>' + str(add_space(10)) + '(<valor>)', tnr_14_000); paragraphs.append(paragraph)

    document.add_page_break()

    paragraph = document.add_paragraph('', tnr_433_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('Valor total dos equipamentos: R$ ' + '<valortotal>', tnr_12_101); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_36_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph(str(add_space(98)) + 'Hortolândia, ' + '<datahoje>', tnr_12_100); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('', tnr_38_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('ANDAIMES PIRÂMIDE IND COM LTDA' + str(add_space(20)) + '<nome>', tnr_12_100); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('CNPJ  56.922.263/0001-02' + str(add_space(50)) + '<x>' + str(add_space(1)) + '<xNumero>', tnr_12_100); paragraphs.append(paragraph)
    paragraph = document.add_paragraph('IE  748.009.175.117' + str(add_space(62)) + '<y>' + str(add_space(1)) + '<yNumero>', tnr_12_100); paragraphs.append(paragraph)
    
    # Remover espaçamento #
    for p in paragraphs:
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = 1
        p.paragraph_format.space_after = 1

    # Salvar documento #
    document.save('TEMPLATE_CONTRATO_PYTHON.docx')   




GERAR_DOCUMENTO()