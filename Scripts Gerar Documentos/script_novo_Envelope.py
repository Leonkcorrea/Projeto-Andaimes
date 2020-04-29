import os
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm
import sqlite3

connection = sqlite3.connect('databaseContatos.db')
c = connection.cursor()

def NOME():
    return 'retornar nome do contrato do banco de dados'

def ENDERECO():
    return 'retornar endereço do contrato do banco de dados'

def BAIRRO():
    return 'retornar bairro do contrato do banco de dados'

def CIDADE():
    return 'retornar cidade do contrato do banco de dados'

def ESTADO():
    return 'retornar estado do contrato do banco de dados'

def CEP():
    return 'retornar cep do contrato do banco de dados'

def GERAR_ENVELOPE(codigo_contrato):
    # Defenindo variáveis #
    document = docx.Document()
    styles = document.styles
    sections = document.sections

    # Configurando a página # -> # Convertendo inch para cm #
    for section in sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(3.55 / 2.54)
        section.left_margin = Inches(3.00 / 2.54)
        section.right_margin = Inches(3.00 / 2.54)
        section.bottom_margin = Inches(2.50 / 2.54)
    
    # Definindo estilos # -> # [fontName_fontSize_bold/italic/underline] #
    cal_14_000 = styles.add_style('CAL_14_000', WD_STYLE_TYPE.PARAGRAPH)

    cal_14_000.font.name = 'Calibri'
    cal_14_000.font.size = Pt(int(cal_14_000.name[4:-4]))
    if(cal_14_000.name[-3] == '1'): cal_14_000.font.bold = True
    if(cal_14_000.name[-2] == '1'): cal_14_000.font.italic = True
    if(cal_14_000.name[-1] == '1'): cal_14_000.font.underline = True

    # Escrevendo no documento #

    paragraphs = []

    paragraph = document.add_paragraph(('<nome>').upper(), cal_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<endereço>', cal_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<cidade>' + ' - ' + '<estado>', cal_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<endereço>', cal_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<A/C  Contas a Pagar>', cal_14_000); paragraphs.append(paragraph)

    paragraph = document.add_paragraph('<cep>', cal_14_000); paragraphs.append(paragraph)

    # Remover espaçamento #
    for p in paragraphs:
        p.paragraph_format.line_spacing = 1.15

    # Salvar documento #
    document.save('TEMPLATE_ENVELOPE_PYTHON.docx') 
    

GERAR_ENVELOPE(0)